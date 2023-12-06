import docx
import openpyxl
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def setFNT(run):
    run.font.name = "標楷體"
    run.font.size = Pt(18)
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')
    run.bold = True

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

num = int(input("請輸入審查委員數:"))
crs = int(input("請輸入課程數:"))
lea = int(input("請選擇聯盟\n1.健康 2.環境 3.終端(若是健康聯盟請輸入1，以此類推):"))

# excel 1-4 課程大表
full = openpyxl.load_workbook('1-4.xlsx') # source 需修改檔名且放在同資料夾!!
shtfull = full.worksheets[lea-1]

# excel 1-2 from google form審查意見
gf = openpyxl.load_workbook('1-2.xlsx')
shtgf = gf.worksheets[0]

for re in range(num): 
    start = True
    for v in range(crs):
        # 從課程大表取資訊
        pln_num = str(shtfull.cell(row=v+4,column=2).value)[0:4] # 取計畫編號
        crs_num = shtfull.cell(row=v+4,column=2).value # 取課程編號
        crs_nam = shtfull.cell(row=v+4,column=8).value # 取課程名稱
        key_mod = shtfull.cell(row=v+4,column=15).value # 取重點模組
        crs_hos = shtfull.cell(row=v+4,column=7).value # 取課程主持人(教師)
        sch = shtfull.cell(row=v+4,column=3).value # 取學校
        dpt = shtfull.cell(row=v+4,column=6).value # 取系所

        re_nam = shtgf.cell(row=re+2,column=3).value[0:3] # 審查委員名稱

        if str(shtfull.cell(row=v+3,column=2).value)[0:4] != str(pln_num):
            chk_same = False # 確認此課程是否跟上個處理的課程同計畫
        else: chk_same = True

        # 重讀乾淨的docx 並寫入後 另存新檔
        if chk_same==False or start==True:
            doc = docx.Document('1-3.docx')
            tbc = 0
            start=False
        else:
            doc = docx.Document(str(pln_num)+str(re_nam)+'.docx')
        tb = doc.tables[tbc]
        tbc+=1

        # 綜合評分
        tmp = shtgf.cell(row=re+2,column=10+v*7).value
        if float(int(tmp))==tmp: tmp=int(tmp)
        score = str(tmp)
        run = tb.rows[11].cells[0].paragraphs[0].add_run(score)
        run.font.name = "標楷體"
        run.font_size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')

        # 評分說明
        ans = ''
        if score=='4': ans+='■特優(4分)\n'
        else: ans+='□特優(4分)\n'
        if score=='3': ans+='■優  (3分)\n'
        else: ans+='□優  (3分)\n'
        if score=='2': ans+='■良  (2分)\n'
        else: ans+='□良  (2分)\n'
        if score=='1': ans+='■差  (1分)\n'
        else: ans+='□差  (1分)\n'
        if score=='4' or score=='3' or score=='2' or score=='1': ans+='□其他'
        else: ans+='■其他('+score+'分)'
        tb.rows[11].cells[5].text = ans

        # 設定寫入之字型及大小
        doc.styles['Normal'].font.name = "標楷體"
        doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')

        # 獲取google表單回覆
        for k in range(4): # 四項審查重點
            result = shtgf.cell(row=re+2,column=4+k+v*7).value[0] #只取答案之第一個字
            if result=='特':
                tb.rows[4+k].cells[5].text = '■特優'
                tb.rows[4+k].cells[5].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            elif result=='優':
                tb.rows[4+k].cells[6].text = '■優'
                tb.rows[4+k].cells[6].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            elif result=='良':
                tb.rows[4+k].cells[7].text = '■良'
                tb.rows[4+k].cells[7].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            else:
                tb.rows[4+k].cells[8].text = '■差'
                tb.rows[4+k].cells[8].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

        # 審查意見
        opin_fir = shtgf.cell(row=re+2,column=8+v*7).value #選取部分
        if opin_fir[0]=='無': opin_fir=''
        opin_ls = opin_fir.split(', ')
        opin_fir=''
        for o in opin_ls:
            opin_fir = opin_fir + o + '\n'
        opin_sec = str(shtgf.cell(row=re+2,column=9+v*7).value) #打字部分
        tb.rows[9].cells[0].text = opin_fir + '\n' + opin_sec

        tb.rows[0].cells[7].text = crs_num # 增加課程編號
        tb.rows[0].cells[3].text = crs_nam # 增加課程名稱
        tb.rows[1].cells[3].text = key_mod # 增加重點模組
        tb.rows[2].cells[3].text = crs_hos # 增加課程主持人
        tb.rows[2].cells[7].text = sch + '/' + dpt # 增加服務單位

        # 以計畫編號及委員姓名存檔
        if str(shtfull.cell(row=v+5,column=2).value)[0:4] != str(pln_num):
            # 增加首頁
            delete_paragraph(doc.paragraphs[11])
            doc.paragraphs[11].insert_paragraph_before()
            rn = doc.paragraphs[11].add_run('\t'*6+'審查委員： '+re_nam) 
            setFNT(rn)
            delete_paragraph(doc.paragraphs[14])
            doc.paragraphs[14].insert_paragraph_before()
            rn = doc.paragraphs[14].add_run('\t'*6+'計畫編號： '+str(pln_num)+'('+str(tbc)+'門課程)')
            setFNT(rn)
            delete_paragraph(doc.paragraphs[15])
            doc.paragraphs[15].insert_paragraph_before()
            rn = doc.paragraphs[15].add_run('\t'*6+'申請學校： '+sch)
            setFNT(rn)
            delete_paragraph(doc.paragraphs[16])
            doc.paragraphs[16].insert_paragraph_before()
            rn = doc.paragraphs[16].add_run('\t'*6+'申請系所： '+dpt)
            setFNT(rn)

            #for pa in range(11,17): doc.paragraphs[pa].alignment = 1

            # 刪除多餘頁面
            # for d in range(tbc-1, 2, 1):
            #     delete_paragraph(doc.tables[d+1])
        
        # if not os.path.exists(str(re_nam)+'審查意見'): os.mkdir(str(re_nam)+'審查意見')
        doc.save(str(pln_num)+str(re_nam)+'.docx')
