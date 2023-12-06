import docx
import openpyxl
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def setFNT(run):
    run.font.name = "標楷體"
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')

num = int(input("請輸入審查委員數:"))
crs = int(input("請輸入課程數:"))
#mp = int(input("請輸入審查重點數"))  # 用來決定google表單回覆該選的column
lea = int(input("請選擇聯盟\n1.健康 2.環境 3.終端(若是健康聯盟請輸入1，以此類推):"))

# 審查委員list，須根據每次審查作調整
comite = ['洪士灝', '馬席彬', '黃世旭', '吳文慶'] 

# excel 1-4 課程大表
full = openpyxl.load_workbook('1-4.xlsx') # source 需修改檔名且放在同資料夾!!
shtfull = full.worksheets[lea-1]

# excel 1-22 from google form審查意見
gf = openpyxl.load_workbook('1-22.xlsx')
shtgf = gf.worksheets[0]

skip = 0
for v in range(crs):
    # 從課程大表取資訊
    pln_num = str(shtfull.cell(row=v+4,column=2).value)[0:4] # 取計畫編號
    crs_num = shtfull.cell(row=v+4,column=2).value # 取課程編號
    crs_nam = shtfull.cell(row=v+4,column=8).value # 取課程名稱
    key_mod = shtfull.cell(row=v+4,column=15).value # 取重點模組
    crs_hos = shtfull.cell(row=v+4,column=7).value # 取課程主持人(教師)
    sch = shtfull.cell(row=v+4,column=3).value # 取學校
    dpt = shtfull.cell(row=v+4,column=6).value # 取系所
    smest = shtfull.cell(row=v+4,column=10).value # 取開課學期

    # 選擇上下學期之寫入對象 (目標要填寫的docx檔)
    if smest[-1]=='2':
        doc = docx.Document('1-6.docx') # for 下學期
    else:
        skip+=1
        continue

    tb = doc.tables[0]

    for rev in range(num):
        # google表單的回應取值對應column， 可能因審查重點變多變少而須調整

        rev_nam = shtgf.cell(row=rev+2,column=3).value[0:3] # 審查委員名稱
        idx = comite.index(str(rev_nam))+1 # 判斷為委員幾(ex:委員1、委員4)

        # 設定寫入之字型及大小
        doc.styles['Normal'].font.name = "Times new roman"
        doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')
        
        # 審查意見
        point = 1
        opin_fir = shtgf.cell(row=rev+2,column=7+(v-skip)*6).value #複選題選取部分
        if opin_fir[0]=='無': opin_fir=''
        opin_ls = opin_fir.split(', ')
        opin_fir=''
        for o in opin_ls:
            if o != '' and o[0] != '無':
                opin_fir = opin_fir + str(point)+'. ' + o + '。\n'
                point+=1
        # 用換行以及句號 分點呈現
        opin_sec = re.split("[\n|。]", str(shtgf.cell(row=rev+2,column=8+(v-skip)*6).value)) #打字部分
        opin_final = '委員'+str(idx)+ '\n' +opin_fir
        for sp in opin_sec:
            if sp != '':
                opin_final = opin_final + str(point) + '. ' + sp + '。\n'
                point+=1
        tb.rows[idx].cells[0].text = opin_final.strip()

    lines = 3 # 根據docx的空行數等等 進行修改
    # 填寫 docx 上面資訊
    rn = doc.paragraphs[0+lines].add_run(str(crs_num)) #課程編號
    setFNT(rn)
    rn = doc.paragraphs[1+lines].add_run(sch+'/'+dpt) #學校/系所
    setFNT(rn)
    rn = doc.paragraphs[2+lines].add_run(crs_nam) #課程名稱
    setFNT(rn)
    rn = doc.paragraphs[3+lines].add_run(crs_hos) #課程教師
    setFNT(rn)

    # 以計畫編號存檔
    doc.save(str(crs_num)+'課程-'+crs_hos+'老師'+smest+'課程審查意見回覆.docx')
