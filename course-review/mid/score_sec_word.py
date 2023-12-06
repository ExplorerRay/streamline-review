import docx
import openpyxl
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def setFNT(run):
    run.font.name = "標楷體"
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')
    run.bold = True

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

num = int(input("請輸入審查委員數:"))
crs = int(input("請輸入課程數:"))
lea = int(input("請選擇聯盟\n1.健康 2.環境 3.終端(若是健康聯盟請輸入1，以此類推):"))

# excel 1-4 課程大表
full = openpyxl.load_workbook('1-4.xlsx') # source 需修改檔名且放在同資料夾!!
shtfull = full.worksheets[lea-1]

# excel 1-22 from google form審查意見 下學期
gf = openpyxl.load_workbook('1-22.xlsx')
shtgf = gf.worksheets[0]

skip = 0
for rev in range(num): 
    doc = docx.Document('3-2.docx')

    for v in range(crs):
        # 從課程大表取資訊
        pln_num = str(shtfull.cell(row=v+4,column=2).value)[0:4] # 取計畫編號
        crs_num = shtfull.cell(row=v+4,column=2).value # 取課程編號
        crs_nam = shtfull.cell(row=v+4,column=8).value # 取課程名稱
        key_mod = shtfull.cell(row=v+4,column=15).value # 取重點模組
        crs_hos = shtfull.cell(row=v+4,column=7).value # 取課程主持人(教師)
        sch = shtfull.cell(row=v+4,column=3).value # 取學校
        dpt = shtfull.cell(row=v+4,column=6).value # 取系所
        smest = shtfull.cell(row=v+4,column=10).value # 取開課學期

        if smest[-1]=='1':
            skip+=1
            continue

        re_nam = shtgf.cell(row=rev+2,column=3).value[0:3] # 審查委員名稱

        tb = doc.tables[0]

        tb.rows[0].cells[0].text = ''
        run = tb.rows[0].cells[0].paragraphs[0].add_run('智慧終端裝置晶片系統與應用聯盟')
        tb.rows[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        setFNT(run)

        # 填入基本資訊
        tb.rows[2+v-skip].cells[0].text = crs_num
        tb.rows[2+v-skip].cells[1].text = sch+'\n'+dpt
        tb.rows[2+v-skip].cells[2].text = crs_nam
        tb.rows[2+v-skip].cells[3].text = crs_hos

        # 綜合評分
        score = shtgf.cell(row=rev+2,column=9+(v-skip)*6).value
        if score>4: score=4
        if isinstance(score, (int, float)):
            score = int(score)
        score = str(score)
        tb.rows[2+v-skip].cells[4].text = score
        

        # 設定寫入之字型及大小
        doc.styles['Normal'].font.name = "標楷體"
        doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')

        # 審查意見
        point = 1
        opin_fir = shtgf.cell(row=rev+2,column=7+(v-skip)*6).value #複選題選取部分
        if opin_fir[0]=='無': opin_fir=''
        opin_ls = opin_fir.split(', ')
        opin_fir=''
        for o in opin_ls:
            if o != '' and o[0] != '無':
                opin_fir = opin_fir + str(point)+'. ' + o + '。\n'
                point+=1
        # 用換行以及句號 分點呈現
        opin_sec = re.split("[\n|。]", str(shtgf.cell(row=rev+2,column=8+(v-skip)*6).value)) #打字部分
        opin_final = opin_fir
        for sp in opin_sec:
            if sp != '':
                opin_final = opin_final + str(point) + '. ' + sp + '。\n'
                point+=1
        tb.rows[2+v-skip].cells[5].text = opin_final.strip()

            #for pa in range(11,17): doc.paragraphs[pa].alignment = 1

            # 刪除多餘頁面
            # for d in range(tbc-1, 2, 1):
            #     delete_paragraph(doc.tables[d+1])
        
        # if not os.path.exists(str(re_nam)+'審查意見'): os.mkdir(str(re_nam)+'審查意見')
    doc.save(str(re_nam)+smest+'.docx')
