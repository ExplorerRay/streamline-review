# try:
import docx
import openpyxl
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def setFNT(run):
    run.font.name = "標楷體"
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')

num = int(input("請輸入審查委員數:"))
md = int(input("請輸入模組數:"))
lea = int(input("請選擇聯盟\n0.健康 1.環境 2.終端(若是健康聯盟請輸入0，以此類推):"))
chk_avd = str(input("是否有利益迴避?(若有請輸入y，否則輸入n):"))

# 審查委員list，須根據每次審查作調整
comite = ['吳安宇', '張孟凡', '鄭國興']

# excel 1-1 分數及意見總表
full = openpyxl.load_workbook('1-1.xlsx') # source 需修改檔名且放在同資料夾!!
shtfull = full.worksheets[lea]

# excel 1-2 from google form審查意見
gf = openpyxl.load_workbook('1-2.xlsx')
shtgf = gf.worksheets[0]

# excel 1 審查委員審查模組表
if chk_avd=='y':
    avd = openpyxl.load_workbook('1.xlsx')
    shtavd = avd.worksheets[0]

for v in range(md):
    # 從分數意見總表取資訊
    mod_num = str(shtfull.cell(row=v+4,column=1).value) # 取模組編號
    mod_nam = shtfull.cell(row=v+4,column=5).value # 取模組名稱
    mod_hos = shtfull.cell(row=v+4,column=4).value # 取模組主持人(教師)
    sch = (shtfull.cell(row=v+4,column=2).value).strip() # 取學校
    dpt = shtfull.cell(row=v+4,column=3).value # 取系所

    # 重讀乾淨的docx 並寫入後 另存新檔
    doc = docx.Document('1-5.docx')

    tb = doc.tables[0]

    for rev in range(num):

        rev_nam = shtgf.cell(row=rev+2,column=3).value[0:3] # 審查委員名稱
        idx = comite.index(str(rev_nam))+1 # 判斷為委員幾(ex:委員1、委員4)

        if chk_avd=='y':
            if str(shtavd.cell(row=v+2,column=idx+1).value) == 'x': continue

        # 設定寫入之字型及大小
        doc.styles['Normal'].font.name = "Times new roman"
        doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')
        
        # 審查意見
        point = 1
        opin_fir = shtgf.cell(row=rev+2,column=8+v*7).value #複選題選取部分
        if opin_fir[0]=='無': opin_fir=''
        opin_ls = opin_fir.split(', ')
        opin_fir=''
        for o in opin_ls:
            if o != '' and o[0] != '無':
                opin_fir = opin_fir + str(point)+'. ' + o + '。\n'
                point+=1
        # 用換行以及句號 分點呈現
        opin_sec = re.split("[\n|。]", str(shtgf.cell(row=rev+2,column=9+v*7).value)) #打字部分
        opin_final = '委員'+str(idx)+ '\n' +opin_fir
        for sp in opin_sec:
            if sp != '':
                opin_final = opin_final + str(point) + '. ' + sp + '。\n'
                point+=1
        tb.rows[idx].cells[0].text = opin_final.strip()

    lines = 2 # 根據1-5.docx的空行數等等 進行修改
    # 填寫1-5上面資訊
    rn = doc.paragraphs[0+lines].add_run(str(mod_num)) #模組編號
    setFNT(rn)
    rn = doc.paragraphs[1+lines].add_run(sch+'/'+dpt) #學校/系所
    setFNT(rn)
    rn = doc.paragraphs[2+lines].add_run(mod_nam) #模組名稱
    setFNT(rn)
    rn = doc.paragraphs[3+lines].add_run(mod_hos) #模組教師
    setFNT(rn)

    # 以計畫編號存檔
    doc.save(str(mod_num)+' '+sch+' '+mod_hos+'_'+mod_nam+'_審查意見回覆.docx')
# except Exception as e:
#     print(e)
#     a=input()
