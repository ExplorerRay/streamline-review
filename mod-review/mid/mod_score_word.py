import docx
import openpyxl
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def setFNT(run):
    run.font.name = "標楷體"
    run.font.size = Pt(18)
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')
    run.bold = True

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

num = int(input("請輸入審查委員數:"))
md = int(input("請輸入模組數:"))
lea = int(input("請選擇聯盟\n0.健康 1.環境 2.終端(若是健康聯盟請輸入0，以此類推):"))
chk_avd = str(input("是否有利益迴避?(若有請輸入y，否則輸入n):"))

# 審查委員list，須根據每次審查作調整
comite = ['吳安宇', '張孟凡', '鄭國興']

# excel 1-1 分數及意見總表
full = openpyxl.load_workbook('1-1.xlsx') # source 需修改檔名且放在同資料夾!!
shtfull = full.worksheets[lea]

# excel 1-2 from google form審查意見
gf = openpyxl.load_workbook('1-2.xlsx')
shtgf = gf.worksheets[0]

# excel 1 審查委員審查模組表
if chk_avd=='y':
    avd = openpyxl.load_workbook('1.xlsx')
    shtavd = avd.worksheets[0]

for rev in range(num): 
    for v in range(md):
        # 從分數及意見總表取資訊
        mod_num = str(shtfull.cell(row=v+4,column=1).value) # 取模組編號
        mod_nam = shtfull.cell(row=v+4,column=5).value # 取模組名稱
        mod_hos = shtfull.cell(row=v+4,column=4).value # 取模組主持人(教師)
        sch = (shtfull.cell(row=v+4,column=2).value).strip() # 取學校
        dpt = shtfull.cell(row=v+4,column=3).value # 取系所

        rev_nam = shtgf.cell(row=rev+2,column=3).value[0:3] # 審查委員名稱
        idx = comite.index(str(rev_nam))+1 # 判斷為委員幾(ex:委員1、委員4)

        if chk_avd=='y':
            if str(shtavd.cell(row=v+2,column=idx+1).value) == 'x': continue

        # 重讀乾淨的docx 並寫入後 另存新檔
        doc = docx.Document('1-3.docx')

        tb = doc.tables[0]

        # 綜合評分
        score = int(shtgf.cell(row=rev+2,column=10+v*7).value)

        # 評等
        ans = ''
        for s in range(10,0,-1):
            if score==s:
                ans+='■'
            else:
                ans+='□'
            ans+=str(s)
            ans+='  '
        ans+='\n'
        ans+='註：(10:極優 9:優 8:良 7:尚可 6:可 5:普通 4:略差 3:差 2:極差 1:劣)'
        tb.rows[9].cells[5].text = ans

        # 設定寫入之字型及大小
        doc.styles['Normal'].font.name = "Times new roman"
        doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')

        # 獲取google表單回覆
        for k in range(4): # 四項審查重點
            result = shtgf.cell(row=rev+2,column=4+k+v*7).value[0] #只取答案之第一個字
            if result=='優':
                tb.rows[3+k].cells[5].text = '■優'
                tb.rows[3+k].cells[5].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            elif result=='佳':
                tb.rows[3+k].cells[6].text = '■佳'
                tb.rows[3+k].cells[6].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            elif result=='尚':
                tb.rows[3+k].cells[8].text = '■尚可'
                tb.rows[3+k].cells[8].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            else:
                tb.rows[3+k].cells[9].text = '■不佳'
                tb.rows[3+k].cells[9].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

        # 審查意見
        point = 1
        opin_fir = shtgf.cell(row=rev+2,column=8+v*7).value #複選題選取部分
        if opin_fir[0]=='無': opin_fir=''
        opin_ls = opin_fir.split(', ')
        opin_fir=''
        for o in opin_ls:
            if o != '' and o[0] != '無':
                opin_fir = opin_fir + str(point)+'. ' + o + '。\n'
                point+=1
        # 用換行以及句號 分點呈現
        opin_sec = re.split("[\n|。]", str(shtgf.cell(row=rev+2,column=9+v*7).value)) #打字部分
        opin_final = opin_fir
        for sp in opin_sec:
            if sp != '':
                opin_final = opin_final + str(point) + '. ' + sp + '。\n'
                point+=1
                
        tb.rows[8].cells[0].text = opin_final.strip()

        tb.rows[0].cells[8].text = mod_num # 增加模組編號
        tb.rows[0].cells[3].text = mod_nam # 增加模組名稱
        tb.rows[1].cells[3].text = mod_hos # 增加模組主持人
        tb.rows[1].cells[8].text = sch + '/' + dpt # 增加服務單位

        if not os.path.exists(str(rev_nam)+'委員評分表'): os.mkdir(str(rev_nam)+'委員評分表')
        doc.save(str(rev_nam)+'委員評分表/'+str(mod_num)+' '+str(rev_nam)+'委員評分表.docx')
