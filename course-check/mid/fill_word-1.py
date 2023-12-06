import docx
import openpyxl

crs = int(input("請輸入總課程數:"))

rs = openpyxl.load_workbook('result1.xlsx')
shtrs = rs.worksheets[0]

doc = docx.Document('1-31.docx')
tb = doc.tables[0]
# 共9 columns, 13 rows
tb.rows[7].cells[1].paragraphs[0].add_run('(詳情請見附件)')

for c in range(crs):
    # 在result有紀錄才填入word
    if shtrs.cell(row=5+c*3, column=18).value==None:
        continue
    
    # 經費執行情形
    tb.rows[2].cells[3].text = str(shtrs.cell(row=5+c*3, column=9).value) #人事
    tb.rows[2].cells[4].text = str(shtrs.cell(row=5+c*3, column=10).value) #業務
    tb.rows[2].cells[7].text = str(shtrs.cell(row=5+c*3, column=11).value) #設備
    tb.rows[2].cells[8].text = str(shtrs.cell(row=5+c*3, column=12).value) #小計

    tb.rows[3].cells[3].text = str(shtrs.cell(row=6+c*3, column=9).value) #人事
    tb.rows[3].cells[4].text = str(shtrs.cell(row=6+c*3, column=10).value) #業務
    tb.rows[3].cells[7].text = str(shtrs.cell(row=6+c*3, column=11).value) #設備
    tb.rows[3].cells[8].text = str(shtrs.cell(row=6+c*3, column=12).value) #小計

    tb.rows[4].cells[3].text = str(shtrs.cell(row=7+c*3, column=9).value) #執行率

    # 教學設備採購進度
    tb.rows[6].cells[2].text = shtrs.cell(row=5+c*3, column=13).value
    tb.rows[6].cells[6].text = shtrs.cell(row=5+c*3, column=14).value

    # 課程模組結合使用情形
    cm=''
    for cnt in range(3):
        if shtrs.cell(row=5+c*3+cnt, column=15).value==None:
            break
        cm = cm + shtrs.cell(row=5+c*3+cnt, column=15).value + ' '
        cm = cm + shtrs.cell(row=5+c*3+cnt, column=16).value + ' '
        cm = cm + shtrs.cell(row=5+c*3+cnt, column=17).value + '\n'
    tb.rows[7].cells[2].text = cm

    # 課程開授成效
    tb.rows[9].cells[2].text = shtrs.cell(row=5+c*3, column=18).value
    tb.rows[9].cells[6].text = shtrs.cell(row=5+c*3, column=19).value

    # 參與聯盟活動競賽情形
    tb.rows[11].cells[2].text = shtrs.cell(row=5+c*3, column=20).value
    tb.rows[11].cells[6].text = shtrs.cell(row=5+c*3, column=21).value

    # 業界或校外講師參與教學情形
    tb.rows[12].cells[2].text = shtrs.cell(row=5+c*3, column=22).value

    doc.save(str(shtrs.cell(row=5+c*3, column=7).value)+'期中查核表.docx')

